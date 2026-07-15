import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold

doc = docx.Document()

# 标题
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run('Math-AI 动态分层智能体：核心功能与技术实现清单')
set_font(run, '黑体', 16, bold=True)

features = [
    {
        "title": "1. 脏数据清洗与缺失值插补 (MICE)",
        "tech": "技术实现：在数据预处理阶段引入多重插补算法 (Multiple Imputation by Chained Equations, MICE)。针对真实教育场景中常见的“未交作业”、“漏测”等缺失值情况，MICE 算法能够利用作业日历和单元测试中的其它特征，对缺失分数进行迭代预测和插补，保证后续学情分析的数据连续性和完整性。"
    },
    {
        "title": "2. 学情异常波动检测与基底能力分离 (RPCA)",
        "tech": "技术实现：采用鲁棒主成分分析算法 (Robust Principal Component Analysis, RPCA)，将学生复杂的分数矩阵分解为低秩矩阵 (L矩阵) 和稀疏矩阵 (S矩阵)。其中，L矩阵(L-Score)用于表征学生真实的基底能力，作为AI进行稳定分层的核心依据；S矩阵(S-Score)用于捕捉学生短期内的突发异常波动（如断崖式掉分），触发系统端的“S矩阵异常干预警报”。这种“先RPCA+MICE清洗，再送入AI分层”的架构，大幅提升了系统的抗干扰能力和分析精准度。"
    },
    {
        "title": "3. 多源学情数据上传与纵向合并存储",
        "tech": "技术实现：基于后端 Hono 框架开发数据流转引擎。支持以CSV格式反复上传“作业日历”与“智学网考情”等多源数据。后端采用“学生ID+教师ID”复合主键策略，不仅实现了不同教师账户之间的数据安全隔离，还支持同一账户不同学段上传数据的动态累加与合并。解决了旧数据被覆盖的问题，从而构建出长达3年的学生纵向学习档案。"
    },
    {
        "title": "4. 大模型驱动的四维动态分层与任务推荐",
        "tech": "技术实现：系统深度对接底层大语言模型（LLM）API。通过构建严谨的 Prompt 规则体系，将 RPCA 提取出的高净值 L-Score 喂入 AI，由 AI 严格依据课题设定的最近发展区理论，将学生动态划分为“反思层、表达层、运用层、理解层”。并基于该阶层，自适应生成包含“概念辨析、变式训练、讲题交流、规律总结”的差异化课堂任务与课后作业。"
    },
    {
        "title": "5. 基于 ECharts 的三维全景知识图谱与全屏高斯视图",
        "tech": "技术实现：在前端 React 框架下深度集成 echarts-for-react。将北师大版初中数学七至九年级的所有章节目录预设为树状数据结构。根据用户上传的学情数据，通过递归算法遍历图谱，计算章节平均得分并映射为四级色彩代码（理解层红、运用层紫、表达层蓝、反思层绿）。通过 React Portal 和 TailwindCSS 打造了支持缩放、拖拽的极大视口与高斯模糊背景全屏 Modal，解决全学段大跨度知识点展示拥挤的痛点。"
    },
    {
        "title": "6. 互动式学情追踪浮层 (Tooltip) 与全局色彩联动",
        "tech": "技术实现：自定义 ECharts 的 tooltip.formatter 函数。通过事件监听捕获用户悬停动作，实时计算父节点（章）下所有子节点（节）的分数均值，并动态拼装出包含得分、阶层定性评价及对应颜色的富文本 DOM 浮层。此外，借助 React 状态管理，实现了左侧班级名册头像背景色、知识图谱节点颜色、以及图例指示器三者之间的全局视觉色彩联通。"
    },
    {
        "title": "7. 学情跃迁轨迹与全景分数画像生成",
        "tech": "技术实现：通过前端数据聚合与 ECharts Line/Bar Chart，开发了“学情跃迁轨迹”折线图和“全景分数画像”模块。将学生在不同章节的 L-Score 表现连接为时间序列，绘制出三年的学业进阶曲线。同时将智学网大考成绩与日常单元测验合并，利用坐标轴映射与自定义 MarkLine（标注四维分层警戒线），直观展示学生能力的宏观演变。"
    },
    {
        "title": "8. 跨平台可编辑数学公式的 Word 文件导出",
        "tech": "技术实现：在实现“分层课堂任务推荐”与“分层课后作业推荐”功能时，考虑到一线数学教师的刚需，前端通过拦截大模型流式输出的富文本与 LaTeX/MathML 标签，利用 HTML 转换技术将其封装为微软 Word 兼容的 XML 结构。确保导出的 Word 文档 (.docx) 中的数学公式、方程式和几何符号能够直接被教研员和教师通过 Office 软件进行二次修改和排版打印。"
    }
]

for feature in features:
    p_h = doc.add_paragraph()
    run = p_h.add_run(feature['title'])
    set_font(run, '黑体', 14, bold=True)
    p_h.paragraph_format.space_before = Pt(12)
    p_h.paragraph_format.space_after = Pt(6)
    
    p_b = doc.add_paragraph()
    run = p_b.add_run(feature['tech'])
    set_font(run, '宋体', 12)
    p_b.paragraph_format.line_spacing = 1.5

doc.save('核心功能与技术实现清单.docx')
print("Document generated successfully.")
