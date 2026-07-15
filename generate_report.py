import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold

doc = docx.Document()

# 标题（黑体三号）
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run('人工智能赋能初中数学动态分层智能体的设计与应用')
set_font(run, '黑体', 16, bold=True)

# 简介（宋体小四）
p_intro = doc.add_paragraph()
run = p_intro.add_run('【简介】本案例立足初中数学大班额课堂个性化教学难的痛点，依托最近发展区与支架理论，自主开发了“Math-AI 动态分层智能体”。该智能体创新性地将传统的“成绩静态分层”转化为基于学习支持需求的“动态分层”，构建了“学情识别—动态分层—任务适配—画像更新”的闭环机制。通过多模态数据采集与AI辅助分析，实现对学生（反思层、表达层、运用层、理解层）的精准研判与个性化任务推送，有效重构了人机共生的教学新生态。')
set_font(run, '宋体', 12)

# 关键词（宋体小四）
p_kw = doc.add_paragraph()
run = p_kw.add_run('【关键词】动态分层；人工智能；初中数学')
set_font(run, '宋体', 12)

# 智能体链接（宋体小四）
p_link = doc.add_paragraph()
run = p_link.add_run('【智能体链接】系统开发环境演示链接（本地）：http://localhost:5173/')
set_font(run, '宋体', 12)

def add_heading1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '黑体', 12, bold=True)

def add_heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '宋体', 12)

def add_heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '宋体', 12)

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, '宋体', 12)

# 一、开发背景
add_heading1('一、开发背景')
add_heading2('（一）政策导向与教育数字化需求')
add_body('《教育强国建设规划纲要》与《成都市基础教育人工智能应用场景建设行动方案》均强调推动人工智能与教育教学深度融合。课堂教学正由经验驱动向数据驱动转型。')
add_heading2('（二）传统大班额教学的现实困境')
add_body('在现有教学中，教师难以及时识别并兼顾学生的差异化学情，学生难以获得适配自身最近发展区的学习任务。传统的固定分层多基于分数，忽视了学生学习支持需求的动态变化，亟待引入AI技术赋能精准教学。')

# 二、拟解决的问题
add_heading1('二、拟解决的问题')
add_heading2('（一）学情诊断的滞后性与经验化问题')
add_body('大班额环境下，教师主要依靠作业反馈和阶段测试了解学情，存在明显的滞后性。难以精准识别学生在概念理解、迁移运用、思维表达和反思监控等维度的即时需求，导致教学决策缺乏科学依据。')
add_heading2('（二）“一刀切”任务与个性化需求的矛盾')
add_body('传统课堂统一的教学目标和任务容易导致“学得快的吃不饱、学得慢的跟不上”。个体差异难以得到回应，缺乏面向学生特定学习障碍（如概念混淆、不会迁移等）的定制化学习支架。')
add_heading2('（三）分层教学固化与实施路径匮乏问题')
add_body('现有的分层多依据成绩进行静态的划分，标签化严重，且缺乏动态流转机制。教师受制于课堂管理难度，难以在常态课中有效组织多任务并行的分层活动。')

# 三、核心功能
add_heading1('三、核心功能')
add_body('本智能体紧扣“教、学、评、研”场景，围绕“理解—运用—表达—反思”四维学情框架，开发了三大核心功能：')
add_heading2('（一）多维数据采集与学情动态研判')
add_body('智能体无缝对接“作业日历”与测试数据，构建常态化的数据流。突破传统的分数评价，基于认知诊断模型，将学情精准划分为“反思层（元认知支持）、表达层（思维表达支持）、运用层（迁移运用支持）、理解层（概念理解支持）”。智能体自动生成学生个人及班级的“全景学情画像”与“知识树结构”，实现学情的伴随式诊断。')
add_heading2('（二）AI驱动的动态分层与任务适配')
add_body('依托大语言模型（LLM）与专属Prompt规则体系，智能体根据学生的动态阶层，自适应调用“动态分层学习任务资源库”。对于理解层提供“概念辨析”支架，运用层提供“变式训练”，表达层提供“讲题交流”，反思层提供“规律总结”。真正实现“学情表现—学习支持—学习任务”的智能匹配与精准推送，且学生层级允许在学习过程中动态向上跃迁或向下调整。')
add_heading2('（三）智能辅助决策与多端协同反馈')
add_body('系统提供“AI辅助分析+教师专业决策”的协同页面。在教师端，提供教学预警（如S矩阵异常干预警报）和微观学情跃迁轨迹；在学生端，可视化呈现作业日历与改错情况。通过“课前画像—课中进阶—课后更新”的逻辑，辅助教研员和教师制定科学的区域或班级教学策略。')

# 四、应用成效
add_heading1('四、应用成效')
add_heading2('（一）学生层面：激发潜能，提升学习效能')
add_body('通过“支持分层”替代“能力分层”，弱化了传统分层的标签效应。学生在最适合自己的最近发展区内获得定制化任务（如基础薄弱生获得概念支架，优生获得拓展挑战），有效缓解了数学学习焦虑，提升了自主探究能力。')
add_heading2('（二）教师层面：数据赋能，实现精准教研')
add_body('智能体为教师提供直观的三维知识图谱与动态折线图，极大降低了学情分析的门槛。教师能够一目了然地定位班级共性薄弱点与个体异常波动，实现从“凭经验备课”向“依数据决策”的跨越。')
add_heading2('（三）学校与区域层面：沉淀资源，形成实践范式')
add_body('沉淀了高质量的学习任务资源库，形成了可复制、可推广的大班额分层教学实施路径，为区域教育数字化转型提供了具有本地特色的生动样本。')

# 五、创新性
add_heading1('五、创新性')
add_heading2('（一）理论创新：首创“学习支持需求导向”的动态分层模式')
add_body('打破传统以分数为单一指标的静态能力分层，以维果茨基最近发展区和支架理论为基础，将分层依据转变为“学生当前最需要获得的学习支持类型”（理解、运用、表达、反思），增强了分层的教育适切性。')
add_heading2('（二）机制创新：构建“学情—分层—任务”的教学闭环')
add_body('将生成式AI深度融入“课前研判、课中进阶、课后更新”全流程，首创学情动态流转机制。学生层级不再一成不变，而是随着任务完成情况实时进阶，形成了动态持续发展的育人新生态。')
add_heading2('（三）技术创新：AI大模型与本土教研的深度融合')
add_body('专为初中数学场景定制，将抽象的教学经验转化为AI可执行的规则库和知识图谱。开发了具有微观指标预警（如异常得分预警）和可视化全景画像的智能体原型，实现了人机协同的高效教研。')

doc.save('智能体设计与应用报告.docx')
print("Document saved successfully.")
